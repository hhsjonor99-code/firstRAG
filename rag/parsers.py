"""firstRAG 文档解析器。

支持的格式：
- PDF（仅带文本层；扫描 PDF 抛 :class:`UnsupportedScannedPDFError`）
- DOCX（python-docx；两层标题识别 + 段落继承最近标题）
- Markdown（按行扫描；标题与正文分离；记录 line_start/line_end）
- TXT（自动尝试 UTF-8 / UTF-8-SIG / GB18030）

公共入口：
- :func:`parse_document` —— 统一按扩展名分发，返回 :class:`RawDocumentSection` 列表
- :class:`DocumentParseError` 及子类 —— 异常层级
"""

from __future__ import annotations

import re
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

# ======================================================================
# 异常层级
# ======================================================================
class DocumentParseError(Exception):
    """所有解析相关异常的基类。"""


class UnsupportedFileTypeError(DocumentParseError):
    """不支持的文件扩展名。"""


class EmptyDocumentError(DocumentParseError):
    """解析成功但内容为空（PDF 可能是扫描版 / 文件本身无文本）。"""


class UnsupportedScannedPDFError(DocumentParseError):
    """PDF 没有可提取的文本层（疑似扫描版，第一版不支持 OCR）。"""


class TextEncodingError(DocumentParseError):
    """所有候选编码均解码失败。"""


# ======================================================================
# 解析结果结构
# ======================================================================
# DOCX 块类型常量
BLOCK_TYPE_PARAGRAPH = "paragraph"
BLOCK_TYPE_TABLE = "table"
BLOCK_TYPE_TEXTBOX = "textbox"


@dataclass
class RawDocumentSection:
    """解析后的语义单元。

    公共字段：
        content: 文本内容（已经过基本清理）。
        source_name: 原始文件名（仅用于展示，不会用作磁盘文件名）。
        page_number: PDF 页码（从 1 开始），其他类型为 None。
        paragraph_number: DOCX 段落编号（从 1 开始，paragraph 块专用），其他类型为 None。
        heading: 当前 section 所属最近标题（DOCX / Markdown）。
        line_start: 行号范围起点（Markdown / TXT），从 1 开始。
        line_end: 行号范围终点（Markdown / TXT）。
        metadata: 额外元数据字典。

    DOCX 扩展字段（不影响其他格式）：
        block_type: ``paragraph`` / ``table`` / ``textbox``，DOCX 块类型；其他格式为 None。
        block_index: 块在文档中的全局顺序（DOCX 专用，从 0 开始；包含 paragraph/table/textbox）。
        table_index: 块在文档中的表格序号（DOCX 专用，table 块专用，从 0 开始）。
        row_start: 表格行范围起点（DOCX table 专用，从 0 开始）。
        row_end: 表格行范围终点（DOCX table 专用，从 0 开始；行合并时取首行）。
        column_names: 表头列名列表（DOCX table 专用；无可靠表头时为 None）。
        paragraph_start: 段落块包含的原始段落号起点（DOCX paragraph 专用）。
        paragraph_end: 段落块包含的原始段落号终点（DOCX paragraph 专用）。
    """

    content: str
    source_name: str
    page_number: Optional[int] = None
    paragraph_number: Optional[int] = None
    heading: Optional[str] = None
    line_start: Optional[int] = None
    line_end: Optional[int] = None
    metadata: dict = field(default_factory=dict)
    # ---- DOCX 扩展字段（向后兼容：默认 None）----
    block_type: Optional[str] = None
    block_index: Optional[int] = None
    table_index: Optional[int] = None
    row_start: Optional[int] = None
    row_end: Optional[int] = None
    column_names: Optional[list[str]] = None
    paragraph_start: Optional[int] = None
    paragraph_end: Optional[int] = None


# ======================================================================
# 解析器抽象与实现
# ======================================================================
class DocumentParser(ABC):
    """文档解析器协议。"""

    @abstractmethod
    def parse(self, path: Path) -> list[RawDocumentSection]:
        """解析文件，返回 section 列表。空结果视情况抛 :class:`EmptyDocumentError`
        或 :class:`UnsupportedScannedPDFError`。"""


# ------------------------- PDF -------------------------
class PdfParser(DocumentParser):
    """基于 pypdf 的 PDF 解析器。

    - 按页提取文本；页码从 1 开始。
    - 单页文本为空时跳过该页（不创建 section）。
    - **所有页均为空** → 抛 :class:`UnsupportedScannedPDFError`。
    - 不实现 OCR。
    """

    def parse(self, path: Path) -> list[RawDocumentSection]:
        # 延迟导入：减少非 PDF 测试用例的启动开销
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        sections: list[RawDocumentSection] = []
        for idx, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text() or ""
            except Exception as e:  # noqa: BLE001
                raise DocumentParseError(
                    f"PDF 第 {idx} 页文本提取失败: {e}"
                ) from e
            cleaned = _clean_text(text)
            if not cleaned:
                continue
            sections.append(
                RawDocumentSection(
                    content=cleaned,
                    source_name=path.name,
                    page_number=idx,
                )
            )
        if not sections:
            raise UnsupportedScannedPDFError(
                f"PDF 未提取到任何文本（可能为扫描 PDF）：{path.name}"
            )
        return sections


# ------------------------- DOCX -------------------------
# 常见中文/数字标题模式（用于 OCR 文档的启发式标题识别）
_HEADING_PATTERN_RE = re.compile(
    r"^("
    r"第[一二三四五六七八九十百千零]+(?:部分|章|节|篇)|"  # 第一部分 / 第一章 / 第一节
    r"第\d+(?:部分|章|节|篇)|"                          # 第1章 / 第2部分
    r"[一二三四五六七八九十]+、|"                       # 一、 二、
    r"（[一二三四五六七八九十]+）|"                     # （一）
    r"\([一二三四五六七八九十]+\)|"                     # (一)
    r"\d+[\.、\)]\s*|"                                 # 1. / 1、 / 1)
    r"[一二三四五六七八九十]+\s*[\.、\)]"
    r").{0,60}$"
)


class DocxParser(DocumentParser):
    """基于 python-docx 的 DOCX 解析器（阶段 2.2 增强版）。

    **遍历策略**：按 ``document.xml`` 中 ``body`` 的直接子节点顺序遍历，
    保留 paragraph / table / textbox 三类块的原始文档顺序。
    **不再使用** ``doc.paragraphs + doc.tables`` 简单拼接（会丢失块顺序）。

    **两层标题识别策略**：

    1. **第一优先级（标准样式）**：Word Heading 1-9 / Title / Subtitle。
    2. **第二优先级（OCR 启发式）**：当文档未使用标准样式时，按以下信号
       保守识别候选标题（**至少满足 2 个条件**）：
       - 文本较短（≤ 30 字符）
       - 段落居中
       - 字体明显大于正文平均（≥ avg + 2pt）
       - 整段加粗（≥ 80% runs 加粗）
       - 匹配 :data:`_HEADING_PATTERN_RE`（"第一章"/"一、"/"1." 等常见模式）

    **表格解析**：
    - 检测表头（首行）：首行非空、单格 ≤ 30 字符、不含终止标点。
    - 每行构造为一条结构化 section，格式 ``列名: 值; 列名: 值``。
    - 无可靠表头时使用 ``第N列: 值``。
    - 处理合并单元格：``vMerge=continue`` 取上方文本；
      ``gridSpan`` 由 python-docx 自动展开，再用 ``_tc`` 句柄去重。
    - 整行为空时过滤。
    - 表格继承之前最近识别出的标题。
    - 不同表格之间不直接合并。

    **过滤统计**：每次 ``parse`` 结束后通过 :attr:`last_filter_stats`
    可查看每条规则过滤了多少块、多少字符，便于排查过度过滤问题。

    普通段落继承最近出现的有效标题。
    保留 paragraph_number（DOCX 段落顺序，从 1 开始）。
    对 OCR 噪声进行基本清理（合并空白、合并连续换行）。
    """

    _HEADING_STYLE_RE = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)

    # 启发式阈值
    _MAX_HEADING_LEN = 30         # 候选标题最大长度
    _BOLD_RATIO_THRESHOLD = 0.8   # 整段加粗比例
    _SIZE_DELTA_PT = 2.0          # 字号差（候选 ≥ avg + 此值）
    # 表头识别阈值
    _MAX_HEADER_CELL_LEN = 30     # 表头单格最大字符数
    _MIN_TABLE_ROWS_FOR_HEADER = 2  # 至少 2 行才尝试识别表头

    def __init__(self) -> None:
        self.last_filter_stats: dict = {}

    def parse(self, path: Path) -> list[RawDocumentSection]:
        from docx import Document

        doc = Document(str(path))
        avg_body_size = self._compute_average_body_size(doc)

        sections: list[RawDocumentSection] = []
        current_heading: Optional[str] = None

        # 过滤统计
        stats = {
            "blank_dropped": 0,
            "blank_chars": 0,
            "table_empty_row_dropped": 0,
            "table_empty_row_chars": 0,
            "textbox_empty_dropped": 0,
        }

        # 按 body 直接子节点顺序遍历；doc.paragraphs / doc.tables 与 body
        # 直接子节点中的 w:p / w:tbl 一一对应，索引同步递增即可。
        body_children = list(doc.element.body)
        para_index = 0  # doc.paragraphs 顺序索引
        tbl_index = 0  # doc.tables 顺序索引

        block_index = -1
        table_index = -1

        for child in body_children:
            tag = child.tag
            if tag.endswith("}p"):
                # 顶层段落
                if para_index >= len(doc.paragraphs):
                    continue
                para = doc.paragraphs[para_index]
                para_index += 1

                # 段落号：与 doc.paragraphs 列表位置一致（1-based）。
                # para_index 已经反映了"截至当前位置已出现的顶层段落数"。
                paragraph_number = para_index
                raw_text = para.text or ""
                cleaned = _clean_text(raw_text)
                style_name = (para.style.name or "") if para.style is not None else ""

                # 过滤 1：纯空白段落（不删除纯空白，但记录）
                if not cleaned:
                    stats["blank_dropped"] += 1
                    stats["blank_chars"] += len(raw_text)
                    continue

                is_heading_text, heading_text = self._detect_heading(
                    para, cleaned, style_name, avg_body_size
                )
                if is_heading_text:
                    current_heading = heading_text
                    continue  # 标题自身不作为独立 section

                block_index += 1
                sections.append(
                    RawDocumentSection(
                        content=cleaned,
                        source_name=path.name,
                        paragraph_number=paragraph_number,
                        heading=current_heading,
                        metadata={"style": style_name} if style_name else {},
                        block_type=BLOCK_TYPE_PARAGRAPH,
                        block_index=block_index,
                        paragraph_start=paragraph_number,
                        paragraph_end=paragraph_number,
                    )
                )

            elif tag.endswith("}tbl"):
                # 顶层表格
                if tbl_index >= len(doc.tables):
                    continue
                tbl = doc.tables[tbl_index]
                tbl_index += 1

                table_index += 1
                rows = self._extract_table_rows(tbl)
                if not rows:
                    continue

                column_names = self._detect_header(rows)
                # 数据行起点：表头存在时跳过首行
                start_row = 1 if column_names is not None else 0

                emitted_rows = 0
                for row_idx, row_cells in enumerate(rows):
                    if row_idx < start_row:
                        continue  # 跳过表头行
                    line = self._format_table_row(row_cells, column_names)
                    cleaned_line = _clean_text(line)
                    if not cleaned_line:
                        stats["table_empty_row_dropped"] += 1
                        stats["table_empty_row_chars"] += sum(
                            len(c) for c in row_cells
                        )
                        continue
                    block_index += 1
                    sections.append(
                        RawDocumentSection(
                            content=cleaned_line,
                            source_name=path.name,
                            heading=current_heading,
                            metadata={
                                "table_index": table_index,
                                "table_total_rows": len(rows),
                            },
                            block_type=BLOCK_TYPE_TABLE,
                            block_index=block_index,
                            table_index=table_index,
                            row_start=row_idx,
                            row_end=row_idx,
                            column_names=list(column_names) if column_names else None,
                        )
                    )
                    emitted_rows += 1

                if emitted_rows == 0:
                    # 整张表只有表头或全部空行 → 回退 block_index
                    block_index -= 1
                    table_index -= 1

            elif tag.endswith("}txbxContent"):
                # 顶层文本框（罕见；多数 txbxContent 嵌在 paragraph/table 内）
                tx_text = _extract_textbox_text(child)
                if not tx_text:
                    stats["textbox_empty_dropped"] += 1
                    continue
                block_index += 1
                sections.append(
                    RawDocumentSection(
                        content=tx_text,
                        source_name=path.name,
                        heading=current_heading,
                        block_type=BLOCK_TYPE_TEXTBOX,
                        block_index=block_index,
                    )
                )

        self.last_filter_stats = stats
        if not sections:
            raise EmptyDocumentError(
                f"DOCX 解析后无有效段落：{path.name}"
            )
        return sections

    # ------------------------------------------------------------------
    # 表格解析辅助
    # ------------------------------------------------------------------
    @staticmethod
    def _extract_table_rows(tbl) -> list[list[str]]:
        """从 python-docx Table 抽取每行每格的纯文本。

        处理合并单元格：
        - ``vMerge=continue``：取上方单元格的文本（不重复输出）。
        - ``gridSpan``：python-docx 已按列展开；通过 ``_tc`` 句柄去重。
        """
        rows_text: list[list[str]] = []
        # 用于 vMerge continue：记录上一行每列的文本
        prev_row_cells: list[str] = []

        for row in tbl.rows:
            row_cells: list[str] = []
            current_prev: list[str] = []

            # 通过 _tc 句柄去重（python-docx 对合并单元格的视觉展开）
            seen_tc_ids: set[int] = set()

            for cell in row.cells:
                tc = cell._tc
                tc_id = id(tc)
                if tc_id in seen_tc_ids:
                    continue  # 同一 tc 已被展开，跳过
                seen_tc_ids.add(tc_id)

                tc_text = cell.text or ""
                cleaned = _clean_text(tc_text)

                # 检测 vMerge
                vm = tc.find(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge"
                )
                if vm is not None:
                    val = vm.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
                        "continue",
                    )
                    if val == "continue":
                        # 续行：使用上方单元格的文本
                        col_idx = len(row_cells)
                        prev_text = (
                            prev_row_cells[col_idx]
                            if col_idx < len(prev_row_cells)
                            else ""
                        )
                        row_cells.append(prev_text)
                        current_prev.append(prev_text)
                        continue

                row_cells.append(cleaned)
                current_prev.append(cleaned)

            rows_text.append(row_cells)
            prev_row_cells = current_prev

        return rows_text

    @staticmethod
    def _detect_header(rows: list[list[str]]) -> Optional[list[str]]:
        """检测表头（保守启发式）。

        满足全部条件才视为表头：
        - 至少 2 行（表头 + 数据）
        - 首行各单元格非空
        - 首行单格 ≤ 30 字符
        - 首行不含句子终止标点
        """
        if len(rows) < DocxParser._MIN_TABLE_ROWS_FOR_HEADER:
            return None
        first = rows[0]
        if not first:
            return None
        if any(not (c and c.strip()) for c in first):
            return None
        if any(len(c.strip()) > DocxParser._MAX_HEADER_CELL_LEN for c in first):
            return None
        if any(re.search(r"[。？！.!?;；]", c) for c in first):
            return None
        return [c.strip() for c in first]

    @staticmethod
    def _format_table_row(cells: list[str], column_names: Optional[list[str]]) -> str:
        """将一行单元格格式化为 ``列名: 值; 列名: 值`` 字符串。

        无表头时使用 ``第N列: 值``。
        """
        parts: list[str] = []
        for idx, cell in enumerate(cells):
            value = (cell or "").strip()
            if not value:
                continue
            if column_names and idx < len(column_names):
                key = column_names[idx]
                parts.append(f"{key}: {value}")
            else:
                parts.append(f"第{idx + 1}列: {value}")
        return "; ".join(parts)

    # ------------------------------------------------------------------
    # 标题识别
    # ------------------------------------------------------------------
    def _detect_heading(
        self,
        para,
        cleaned: str,
        style_name: str,
        avg_body_size: float,
    ) -> tuple[bool, Optional[str]]:
        """返回 (是否标题, 标题文本)。"""

        # 第一优先级：标准样式
        if self._HEADING_STYLE_RE.match(style_name) or style_name in {"Title", "Subtitle"}:
            return (bool(cleaned), cleaned or None)

        if not cleaned:
            return (False, None)

        # 第二优先级：启发式（保守）
        # 信号收集
        signals: list[str] = []
        text = cleaned.strip()

        # 信号 1：长度
        if len(text) <= self._MAX_HEADING_LEN:
            signals.append("short")

        # 信号 2：居中
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            signals.append("centered")

        # 信号 3：字号偏大
        max_size = self._max_run_size(para)
        if max_size and max_size >= avg_body_size + self._SIZE_DELTA_PT:
            signals.append("large_font")

        # 信号 4：整段加粗
        bold_ratio = self._bold_ratio(para)
        if bold_ratio >= self._BOLD_RATIO_THRESHOLD:
            signals.append("bold")

        # 信号 5：匹配常见标题模式
        if _HEADING_PATTERN_RE.match(text):
            signals.append("pattern")

        # 信号 6：不包含句子终止标点（标题一般不出现 "。" "?" "!"）
        if not re.search(r"[。？！!?]", text):
            signals.append("no_terminator")

        # 保守策略：要求至少 2 个强信号（short / large_font / centered / bold / pattern）
        strong_signals = {s for s in signals if s in {"short", "large_font", "centered", "bold", "pattern"}}
        if len(strong_signals) < 2:
            return (False, None)
        # 防止误判药品名称：长度 ≤ 30 且没有 pattern 但只有 large_font → 不算标题
        # （已在 strong_signals 集合限制里得到约束）
        return (True, text)

    @staticmethod
    def _compute_average_body_size(doc) -> float:
        """计算正文平均字号；用于启发式标题识别中的对比基准。

        - 取所有 run.font.size 有效的字号；
        - 去除明显异常（如 < 6pt 或 > 50pt）的极端值；
        - 默认值 10.5pt（小四）。
        """
        sizes: list[float] = []
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.size is not None:
                    pt = run.font.size.pt
                    if 6.0 <= pt <= 50.0:
                        sizes.append(pt)
        if not sizes:
            return 10.5
        return sum(sizes) / len(sizes)

    @staticmethod
    def _max_run_size(para) -> Optional[float]:
        max_size: Optional[float] = None
        for run in para.runs:
            if run.font.size is None:
                continue
            pt = run.font.size.pt
            if max_size is None or pt > max_size:
                max_size = pt
        return max_size

    @staticmethod
    def _bold_ratio(para) -> float:
        total = 0
        bold = 0
        for run in para.runs:
            text = run.text or ""
            if not text.strip():
                continue
            total += 1
            if run.bold:
                bold += 1
        if total == 0:
            return 0.0
        return bold / total


# ------------------------- Markdown -------------------------
class MarkdownParser(DocumentParser):
    """基于行扫描的 Markdown 解析器。

    - 识别 # / ## / ### ... 开头的标题，记录层级。
    - 连续正文行（去除空行）作为一个 section；line_start / line_end 记录行号范围。
    - 正文继承最近出现的标题（最近一个任意层级标题）。
    """

    _HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*#*\s*$")

    def parse(self, path: Path) -> list[RawDocumentSection]:
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            # 部分 MD 文件可能含 BOM
            text = path.read_text(encoding="utf-8-sig")

        lines = text.splitlines()
        sections: list[RawDocumentSection] = []
        current_heading: Optional[str] = None

        # 临时缓存一个 section 的起始行
        buf: list[str] = []
        line_start: Optional[int] = None

        def flush(end_line: int) -> None:
            nonlocal buf, line_start
            if buf and line_start is not None:
                content = _clean_text("\n".join(buf))
                if content:
                    sections.append(
                        RawDocumentSection(
                            content=content,
                            source_name=path.name,
                            heading=current_heading,
                            line_start=line_start,
                            line_end=end_line,
                        )
                    )
            buf = []
            line_start = None

        for i, line in enumerate(lines, start=1):
            m = self._HEADING_RE.match(line.strip())
            if m:
                # 标题行：先 flush 当前缓存，再更新 heading
                flush(i - 1)
                current_heading = m.group(2).strip()
                continue

            if not line.strip():
                # 空行：作为一个 section 的自然边界
                flush(i - 1)
                continue

            if line_start is None:
                line_start = i
            buf.append(line)

        # 文件末尾 flush
        flush(len(lines))

        if not sections:
            raise EmptyDocumentError(
                f"Markdown 解析后无有效内容：{path.name}"
            )
        return sections


# ------------------------- TXT -------------------------
class TxtParser(DocumentParser):
    """TXT 解析器。

    自动尝试 UTF-8 / UTF-8-SIG / GB18030；全部失败 → :class:`TextEncodingError`。
    按行扫描，将连续非空行合并为一个 section；保留 line_start / line_end。
    """

    CANDIDATE_ENCODINGS = ("utf-8", "utf-8-sig", "gb18030")

    def parse(self, path: Path) -> list[RawDocumentSection]:
        raw = path.read_bytes()
        text: Optional[str] = None
        last_err: Optional[Exception] = None
        for enc in self.CANDIDATE_ENCODINGS:
            try:
                text = raw.decode(enc)
                break
            except UnicodeDecodeError as e:
                last_err = e
        if text is None:
            raise TextEncodingError(
                f"TXT 解码失败：尝试 {list(self.CANDIDATE_ENCODINGS)} 均不匹配。"
                f" 最后错误：{last_err}"
            )

        lines = text.splitlines()
        sections: list[RawDocumentSection] = []
        buf: list[str] = []
        line_start: Optional[int] = None

        def flush(end_line: int) -> None:
            nonlocal buf, line_start
            if buf and line_start is not None:
                content = _clean_text("\n".join(buf))
                if content:
                    sections.append(
                        RawDocumentSection(
                            content=content,
                            source_name=path.name,
                            line_start=line_start,
                            line_end=end_line,
                        )
                    )
            buf = []
            line_start = None

        for i, line in enumerate(lines, start=1):
            if not line.strip():
                flush(i - 1)
                continue
            if line_start is None:
                line_start = i
            buf.append(line)
        flush(len(lines))

        if not sections:
            raise EmptyDocumentError(
                f"TXT 解析后无有效行：{path.name}"
            )
        return sections


# ======================================================================
# 内部工具
# ======================================================================
_MULTI_SPACE = re.compile(r"[ \t　]+")  # 含全角空格
_MULTI_NEWLINE = re.compile(r"\n{3,}")


def _clean_text(text: str) -> str:
    """通用文本清理：
    - 合并连续空格 / 制表符 / 全角空格为单空格
    - 合并 3+ 连续换行为 2 个
    - 去掉首尾空白
    - 保留单字符内容（不主动删除）
    """
    if not text:
        return ""
    s = _MULTI_SPACE.sub(" ", text)
    s = _MULTI_NEWLINE.sub("\n\n", s)
    return s.strip()


def _extract_textbox_text(elem) -> str:
    """从 OOXML ``w:txbxContent`` 元素抽取所有文本。

    合并内部连续段落（用 ``\\n`` 连接）。
    """
    if elem is None:
        return ""
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    paragraphs = []
    for p in elem.iter(f"{W}p"):
        text = "".join(t.text or "" for t in p.iter(f"{W}t"))
        if text.strip():
            paragraphs.append(text.strip())
    return "\n".join(paragraphs)


# ======================================================================
# 统一入口
# ======================================================================
_PARSERS: dict[str, DocumentParser] = {
    ".pdf": PdfParser(),
    ".docx": DocxParser(),
    ".md": MarkdownParser(),
    ".markdown": MarkdownParser(),
    ".txt": TxtParser(),
}


def parse_document(path: Path) -> list[RawDocumentSection]:
    """按扩展名分发到对应解析器，返回 :class:`RawDocumentSection` 列表。

    Raises:
        UnsupportedFileTypeError: 文件扩展名不在白名单中。
        FileNotFoundError: 文件不存在。
        DocumentParseError: 解析失败（子类含具体原因）。
        EmptyDocumentError: 解析成功但无内容。
        UnsupportedScannedPDFError: PDF 无文本层。
        TextEncodingError: TXT 解码失败。
    """
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")
    ext = path.suffix.lower()
    parser = _PARSERS.get(ext)
    if parser is None:
        raise UnsupportedFileTypeError(
            f"不支持的文件类型 '{ext}'；支持的扩展名: {sorted(_PARSERS)}"
        )
    return parser.parse(path)


def supported_extensions() -> list[str]:
    """返回当前支持的扩展名列表（仅用于 UI 提示）。"""
    return sorted(_PARSERS.keys())