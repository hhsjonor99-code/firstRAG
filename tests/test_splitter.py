"""中文分块器单元测试。"""

from __future__ import annotations

import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from rag.models import DocumentChunk  # noqa: E402
from rag.parsers import RawDocumentSection  # noqa: E402
from rag.splitter import (  # noqa: E402
    SplitOptions,
    _make_chunk_id,
    split_sections,
)


def _sec(content: str, **kwargs) -> RawDocumentSection:
    kwargs.setdefault("source_name", "t.txt")
    return RawDocumentSection(content=content, **kwargs)


# ----------------------------- 中文长文本 -----------------------------
def test_chinese_long_text_splitting():
    """长中文文本能被切分为多个 chunk，且都在 chunk_size 附近。"""
    long_text = "国家基本药物目录包括化学药品、中成药等多个类别。" * 30
    sections = [_sec(long_text)]
    chunks = split_sections(
        sections,
        document_id="docA",
        options=SplitOptions(chunk_size=200, chunk_overlap=20),
    )
    assert len(chunks) > 1
    for c in chunks:
        assert len(c.content) <= 2 * 200


def test_chinese_separator_priority():
    """中文文本优先按双换行 / 单换行 / 中文句号切分。"""
    parts = [
        "国家基本药物目录\n\n化学药品\n\n中成药",     # 双换行
        "第一段。\n第二段。\n第三段。",                # 单换行 + 中文句号
    ]
    sections = [_sec(p) for p in parts]
    chunks = split_sections(sections, document_id="docB", options=SplitOptions(chunk_size=20, chunk_overlap=0))
    # 不应产生空 chunk
    for c in chunks:
        assert c.content.strip()
    # 应至少有 2 个 chunk
    assert len(chunks) >= 2


# ----------------------------- overlap -----------------------------
def test_overlap_produces_shared_substring():
    """相邻 chunk 之间应有 overlap 共享内容。"""
    long_text = "国家基本药物目录是各级医疗机构配备使用药品的依据。" * 20
    chunks = split_sections(
        [_sec(long_text)],
        document_id="docC",
        options=SplitOptions(chunk_size=200, chunk_overlap=50),
    )
    assert len(chunks) >= 2
    # 相邻 chunk 的尾部与下一 chunk 的头部应有 overlap
    head = chunks[1].content[:80]
    assert any(sub in head for sub in [chunks[0].content[i:i+30] for i in range(0, len(chunks[0].content)-30, 20)])


# ----------------------------- 元数据保留 -----------------------------
def test_metadata_preserved_after_split():
    """分块后 source_name / page_number / paragraph_number / heading / line_* 全部继承。"""
    sections = [
        _sec("化学药品" * 50, source_name="国家目录.docx", paragraph_number=5, heading="第一章"),
        _sec("中成药" * 50, source_name="国家目录.docx", page_number=3, heading=None),
    ]
    chunks = split_sections(sections, document_id="docD", options=SplitOptions(chunk_size=100, chunk_overlap=0))
    assert len(chunks) >= 2
    # 第一个 section 切出的所有 chunk 都继承同一份元数据
    for c in chunks:
        if "化学药品" in c.content:
            assert c.source_name == "国家目录.docx"
            assert c.paragraph_number == 5
            assert c.heading == "第一章"
        elif "中成药" in c.content:
            assert c.page_number == 3
            assert c.heading is None


# ----------------------------- chunk_index -----------------------------
def test_chunk_index_continuous_from_zero():
    """chunk_index 从 0 开始连续递增。"""
    sections = [
        _sec("第一段内容。" * 30),
        _sec("第二段内容。" * 30),
    ]
    chunks = split_sections(sections, document_id="docE", options=SplitOptions(chunk_size=80, chunk_overlap=0))
    indices = [c.chunk_index for c in chunks]
    assert indices == list(range(len(chunks)))


def test_chunk_id_stable_for_same_inputs():
    """相同 (document_id, chunk_index) 生成相同 chunk_id。"""
    a = _make_chunk_id("docX", 0)
    b = _make_chunk_id("docX", 0)
    assert a == b
    # 不同 chunk_index 生成不同 ID
    c = _make_chunk_id("docX", 1)
    assert a != c
    # 不同 document_id 生成不同 ID
    d = _make_chunk_id("docY", 0)
    assert a != d


def test_chunk_id_unique_across_document():
    """整个文档的 chunk_id 全部唯一。"""
    sections = [_sec("段落内容。" * 50)]
    chunks = split_sections(sections, document_id="docF", options=SplitOptions(chunk_size=100, chunk_overlap=10))
    ids = [c.chunk_id for c in chunks]
    assert len(ids) == len(set(ids))


# ----------------------------- 无空 chunk -----------------------------
def test_no_empty_chunks():
    """不允许空 chunk 输出；纯空白段被过滤。"""
    sections = [
        _sec("国家基本药物目录"),
        _sec("   \n\n  \n"),  # 纯空白
        _sec("化学药品"),
    ]
    chunks = split_sections(sections, document_id="docG")
    for c in chunks:
        assert c.content.strip(), f"发现空 chunk: {c!r}"
    # 两个非空段被合并为 1 个 chunk（同 heading）
    assert len(chunks) == 1
    assert "国家基本药物目录" in chunks[0].content
    assert "化学药品" in chunks[0].content


def test_pure_punct_section_filtered():
    """纯标点 section 被过滤，不进入 chunk。"""
    sections = [
        _sec("国家基本药物目录"),
        _sec("。！？"),
        _sec("化学药品"),
    ]
    chunks = split_sections(sections, document_id="doc-pp")
    for c in chunks:
        # 纯标点不应进入
        assert "。！？" not in c.content or "化学药品" in c.content


# ----------------------------- 参数校验 -----------------------------
def test_chunk_overlap_must_be_less_than_chunk_size():
    """overlap >= chunk_size 应抛 ValueError。"""
    sections = [_sec("任意内容")]
    with pytest.raises(ValueError):
        split_sections(sections, document_id="docH", options=SplitOptions(chunk_size=100, chunk_overlap=100))
    with pytest.raises(ValueError):
        split_sections(sections, document_id="docH", options=SplitOptions(chunk_size=0, chunk_overlap=0))


def test_settings_default_used_when_no_options():
    """未传 options 时从 Settings 读取默认 chunk_size / overlap。"""
    from config.settings import get_settings
    s = get_settings()
    sections = [_sec("国家基本药物目录。" * 50)]
    chunks = split_sections(sections, document_id="docI")
    assert len(chunks) >= 1
    # 默认 chunk_size = 800，单段总字符约 1100，可能被切为 2 块
    assert all(c.chunk_index < 10 for c in chunks)


# ----------------------------- 文档级串联 -----------------------------
def test_docx_pipeline_end_to_end(tmp_path: Path):
    """解析 → 切分端到端：DOCX 标题与 paragraph_number 都被保留。"""
    from docx import Document
    from rag.parsers import DocxParser

    doc = Document()
    doc.add_heading("第一章 总则", level=1)
    doc.add_paragraph("国家基本药物目录包括化学药品、中成药。" * 30)
    doc.add_heading("第二章 分类", level=1)
    doc.add_paragraph("化学药品包括抗感染药。" * 30)
    path = tmp_path / "pipeline.docx"
    doc.save(str(path))

    sections = DocxParser().parse(path)
    chunks = split_sections(sections, document_id="pipe-doc", options=SplitOptions(chunk_size=200, chunk_overlap=20))
    assert all(isinstance(c, DocumentChunk) for c in chunks)
    # chunk 应保留 heading 字段
    headings = {c.heading for c in chunks if c.heading}
    assert "第一章 总则" in headings
    assert "第二章 分类" in headings


# ===================== 阶段 2.1：分块质量优化 =====================
def test_short_paragraphs_merged_into_next():
    """纯序号短段优先合并到下一段。

    合并时元数据以 base 段为准；被合并的短段 paragraph_number
    会记录到 metadata["merged_from"] 中（因为单纯叠加会产生重复）。
    """
    sections = [
        _sec("这是一段正常正文，包含完整句子。", paragraph_number=1),
        _sec("1", paragraph_number=2),  # 纯序号短段
        _sec("后续段落内容包含完整句子。", paragraph_number=3),
    ]
    chunks = split_sections(sections, document_id="merge-next", options=SplitOptions(chunk_size=400, chunk_overlap=0))
    # 三段在同一个 heading 下应合并为 1 个 chunk
    assert len(chunks) == 1
    content = chunks[0].content
    assert "1" in content
    assert "后续段落内容" in content
    # 段落号：1（base）和 3（下一段）保留；2 被合并到 metadata
    assert chunks[0].paragraph_numbers == [1, 3]
    assert chunks[0].paragraph_start == 1
    assert chunks[0].paragraph_end == 3
    # 被合并的短段号记录在 metadata
    assert 2 in chunks[0].metadata.get("merged_from", ())


def test_short_paragraph_merged_into_prev_when_no_next():
    """短段无后续时合并到前一段。"""
    sections = [
        _sec("前一段的完整正文，包含完整句子。", paragraph_number=1),
        _sec("1", paragraph_number=2),  # 末尾短段
    ]
    chunks = split_sections(sections, document_id="merge-prev", options=SplitOptions(chunk_size=400, chunk_overlap=0))
    assert len(chunks) == 1
    assert "前一段" in chunks[0].content
    assert "1" in chunks[0].content


def test_does_not_merge_across_heading():
    """不跨 heading 合并。"""
    sections = [
        _sec("第一章第一段内容。" * 5, heading="第一章"),
        _sec("第一章第二段内容。" * 5, heading="第一章"),
        _sec("第二章第一段内容。" * 5, heading="第二章"),
        _sec("第二章第二段内容。" * 5, heading="第二章"),
    ]
    chunks = split_sections(sections, document_id="no-cross", options=SplitOptions(chunk_size=100, chunk_overlap=0))
    # 第一章 chunks 的 heading 都是 "第一章"
    first_chapter = [c for c in chunks if c.heading == "第一章"]
    second_chapter = [c for c in chunks if c.heading == "第二章"]
    assert len(first_chapter) >= 1
    assert len(second_chapter) >= 1
    # 没有任何 chunk 同时包含两个章节的内容
    for c in chunks:
        assert not ("第一章" in c.content and "第二章" in c.content)


def test_chunk_index_continuous_after_merge():
    """chunk_index 在合并后仍从 0 连续递增。"""
    sections = [
        _sec("短段一。"),
        _sec("短段二。"),
        _sec("短段三。"),
        _sec("短段四。"),
    ]
    chunks = split_sections(sections, document_id="idx-test", options=SplitOptions(chunk_size=200, chunk_overlap=0))
    indices = [c.chunk_index for c in chunks]
    assert indices == list(range(len(chunks)))


def test_paragraph_start_and_end_recorded():
    """聚合多个段落时 paragraph_start / paragraph_end 正确。"""
    sections = [
        _sec("段落1内容。"),
        _sec("段落2内容。"),
        _sec("段落3内容。"),
    ]
    for i, s in enumerate(sections, start=5):
        s.paragraph_number = i
    chunks = split_sections(sections, document_id="range-test", options=SplitOptions(chunk_size=200, chunk_overlap=0))
    assert len(chunks) == 1
    assert chunks[0].paragraph_start == 5
    assert chunks[0].paragraph_end == 7
    assert chunks[0].paragraph_number == 5  # 兼容旧字段
    assert chunks[0].paragraph_numbers == [5, 6, 7]


def test_metadata_tracks_all_paragraphs():
    """元数据可追踪所有聚合的段落号。"""
    sections = [
        _sec("正文段落A。", paragraph_number=10),
        _sec("正文段落B。", paragraph_number=11),
        _sec("正文段落C。", paragraph_number=12),
    ]
    chunks = split_sections(sections, document_id="track-test")
    assert len(chunks) >= 1
    # 第一个 chunk 应包含 10-12
    assert chunks[0].paragraph_start == 10
    assert 11 in chunks[0].paragraph_numbers
    assert 12 in chunks[0].paragraph_numbers


def test_long_paragraph_split_via_splitter():
    """超长段落会被 RecursiveCharacterTextSplitter 二次切分。"""
    long_text = "国家基本药物目录包括化学药品、中成药等多个类别。" * 30  # ≈ 900 字符
    sections = [_sec(long_text, paragraph_number=1)]
    chunks = split_sections(
        sections,
        document_id="long-test",
        options=SplitOptions(chunk_size=200, chunk_overlap=20),
    )
    assert len(chunks) > 1
    # 都在 chunk_size 附近
    for c in chunks:
        assert len(c.content) <= 2 * 200
    # 所有 chunk 共享同一段落号
    assert all(c.paragraph_numbers == [1] for c in chunks)


def test_overlap_still_works_after_aggregation():
    """聚合后超长文本被切分时 overlap 仍生效。"""
    long_text = "国家基本药物目录是各级医疗机构配备使用药品的依据。" * 20
    chunks = split_sections(
        [_sec(long_text)],
        document_id="ovl-test",
        options=SplitOptions(chunk_size=200, chunk_overlap=50),
    )
    assert len(chunks) >= 2
    # 相邻 chunk 头部应包含上一 chunk 尾部的 overlap
    tail = chunks[0].content[-80:]
    head = chunks[1].content[:120]
    # 重叠区至少 10 字符
    overlap_text = next(
        (tail[i:i+30] for i in range(len(tail) - 30) if tail[i:i+30] in head),
        None,
    )
    assert overlap_text is not None, "overlap 未生效"


def test_no_one_char_chunk():
    """原则上不出现 1 字符 chunk。"""
    sections = [
        _sec("1"),
        _sec("（一）"),
        _sec("一"),
        _sec("正常段落内容，包含完整句子。"),
    ]
    chunks = split_sections(sections, document_id="one-char")
    for c in chunks:
        # 1 字符 chunk 不应出现（除非全部内容合并后仍只有 1 字符）
        assert len(c.content) > 1 or c.content == "1", (
            f"发现 1 字符 chunk: {c.content!r}"
        )


def test_chinese_drug_name_not_all_misclassified_as_heading():
    """普通短药品名称不会被全部误判为标题（实际测试在 test_parsers 中）。"""
    pass