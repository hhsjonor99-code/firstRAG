"""文档解析器单元测试。"""

from __future__ import annotations

import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from rag.parsers import (  # noqa: E402
    DocxParser,
    EmptyDocumentError,
    MarkdownParser,
    PdfParser,
    RawDocumentSection,
    TextEncodingError,
    TxtParser,
    UnsupportedFileTypeError,
    UnsupportedScannedPDFError,
    parse_document,
    supported_extensions,
)


FIXTURES = ROOT / "tests" / "fixtures"


# ----------------------------- DOCX -----------------------------
def _make_docx(tmp_path: Path, title_to_paragraphs: dict[str, list[str]]) -> Path:
    """用 python-docx 在临时目录生成测试 DOCX。"""
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()
    # 设置一些标题样式名称（python-docx 默认就有 Heading 1, Heading 2 等）
    # 直接使用即可
    for heading, paragraphs in title_to_paragraphs.items():
        # 标题
        doc.add_heading(heading, level=1)
        for p in paragraphs:
            doc.add_paragraph(p)
    out = tmp_path / "test.docx"
    doc.save(str(out))
    return out


def test_docx_heading_inheritance(tmp_path: Path):
    """DOCX 段落继承最近出现的标题。"""
    docx = _make_docx(
        tmp_path,
        {
            "第一章 总则": ["第一条  本目录适用于各级医疗机构。", "第二条  内容略。"],
            "第二章 分类": ["化学药品", "中成药"],
        },
    )
    sections = DocxParser().parse(docx)
    # 4 个正文段落
    assert len(sections) == 4
    # 前两个继承"第一章 总则"
    assert sections[0].heading == "第一章 总则"
    assert sections[1].heading == "第一章 总则"
    # 后两个继承"第二章 分类"
    assert sections[2].heading == "第二章 分类"
    assert sections[3].heading == "第二章 分类"


def test_docx_paragraph_numbering(tmp_path: Path):
    """DOCX 段落编号从 1 开始，连续递增。"""
    docx = _make_docx(
        tmp_path,
        {
            "标题A": ["段落1内容", "段落2内容", "段落3内容"],
        },
    )
    sections = DocxParser().parse(docx)
    assert [s.paragraph_number for s in sections] == [2, 3, 4]  # 标题是第 1 段
    # 注意：标题自身不计为 section


def test_docx_keeps_single_char_paragraphs(tmp_path: Path):
    """不删除所有单字符段落（如中文编号 "一" "二"）。

    注：parser 启发式可能把 "一、概述" 识别为标题（pattern + short）；
    此测试用不会触发标题识别的开头段落，仅校验后两个短段保留。
    """
    from docx import Document

    doc = Document()
    doc.add_paragraph("这是一段用于测试的中文正文段落，包含多句话。")
    doc.add_paragraph("甲、内容一")
    doc.add_paragraph("乙、内容二")
    out = tmp_path / "single.docx"
    doc.save(str(out))

    sections = DocxParser().parse(out)
    assert len(sections) == 3
    assert "甲、内容一" in sections[1].content
    assert "乙、内容二" in sections[2].content


def test_docx_cleans_excess_whitespace(tmp_path: Path):
    """OCR 风格的多余空格/换行被清理。"""
    from docx import Document

    doc = Document()
    doc.add_paragraph("国家基本药物目录\n\n\n   第一章   总   则")
    doc.save(str(tmp_path / "ws.docx"))
    sections = DocxParser().parse(tmp_path / "ws.docx")
    assert len(sections) == 1
    # 多余空白被合并
    assert "\n\n\n" not in sections[0].content
    assert "   " not in sections[0].content


def test_docx_empty_raises(tmp_path: Path):
    """DOCX 没有任何段落 → EmptyDocumentError。"""
    from docx import Document

    doc = Document()
    out = tmp_path / "empty.docx"
    doc.save(str(out))
    with pytest.raises(EmptyDocumentError):
        DocxParser().parse(out)


# ----------------------------- TXT -----------------------------
@pytest.mark.parametrize("encoding", ["utf-8", "utf-8-sig", "gb18030"])
def test_txt_encoding_detection(tmp_path: Path, encoding: str):
    """TXT 自动尝试多种中文编码。"""
    content = "国家基本药物目录\n化学药品\n中成药\n"
    path = tmp_path / "sample.txt"
    path.write_bytes(content.encode(encoding))
    sections = TxtParser().parse(path)
    # 连续非空行合并为一个 section
    assert len(sections) >= 1
    assert "国家基本药物目录" in sections[0].content
    assert "化学药品" in sections[0].content


def test_txt_unsupported_encoding_raises(tmp_path: Path):
    """所有候选编码失败 → TextEncodingError。"""
    path = tmp_path / "bad.txt"
    # 写一个 UTF-16 LE 字节序列（不在候选列表中）
    path.write_bytes(b"\xff\xfeH\x00i\x00")
    with pytest.raises(TextEncodingError):
        TxtParser().parse(path)


def test_txt_records_line_numbers(tmp_path: Path):
    """TXT 保留 line_start / line_end。"""
    path = tmp_path / "lines.txt"
    path.write_text("第一行\n第二行\n第三行\n", encoding="utf-8")
    sections = TxtParser().parse(path)
    # 3 行连续非空 → 1 个 section，line range 1-3
    assert len(sections) == 1
    assert sections[0].line_start == 1
    assert sections[0].line_end == 3


def test_txt_blank_lines_split_sections(tmp_path: Path):
    """空行作为 section 边界。"""
    path = tmp_path / "split.txt"
    path.write_text("第一段第一行\n第一段第二行\n\n第二段内容\n", encoding="utf-8")
    sections = TxtParser().parse(path)
    assert len(sections) == 2
    assert sections[0].line_start == 1 and sections[0].line_end == 2
    assert sections[1].line_start == 4 and sections[1].line_end == 4


def test_txt_empty_raises(tmp_path: Path):
    path = tmp_path / "empty.txt"
    path.write_text("   \n  \n  \n", encoding="utf-8")
    with pytest.raises(EmptyDocumentError):
        TxtParser().parse(path)


# ----------------------------- Markdown -----------------------------
def test_markdown_heading_and_line_numbers(tmp_path: Path):
    md = (
        "# 第一章 概述\n"
        "本章介绍国家基本药物目录。\n"
        "包含化学药品和中成药。\n"
        "\n"
        "## 1.1 化学药品\n"
        "化学药品包括抗感染药等。\n"
        "## 1.2 中成药\n"
        "中成药部分略。\n"
    )
    path = tmp_path / "doc.md"
    path.write_text(md, encoding="utf-8")
    sections = MarkdownParser().parse(path)
    # 3 个非空 section：第一章正文、1.1 正文、1.2 正文
    assert len(sections) == 3
    # 第一个 section 继承 # 第一章
    assert sections[0].heading == "第一章 概述"
    assert sections[0].line_start == 2
    assert sections[0].line_end == 3
    # 第二个 section 继承 ## 1.1
    assert sections[1].heading == "1.1 化学药品"
    # 第三个 section 继承 ## 1.2
    assert sections[2].heading == "1.2 中成药"


def test_markdown_inherits_heading_hierarchy(tmp_path: Path):
    """任意层级标题都会被正文继承。"""
    md = (
        "# 顶层标题\n"
        "顶层正文。\n"
        "## 二级标题\n"
        "二级正文。\n"
    )
    path = tmp_path / "h.md"
    path.write_text(md, encoding="utf-8")
    sections = MarkdownParser().parse(path)
    # 顶层正文继承"顶层标题"，二级正文继承"二级标题"
    assert sections[0].heading == "顶层标题"
    assert sections[1].heading == "二级标题"


def test_markdown_empty_raises(tmp_path: Path):
    path = tmp_path / "empty.md"
    path.write_text("# 标题\n\n\n", encoding="utf-8")
    with pytest.raises(EmptyDocumentError):
        MarkdownParser().parse(path)


# ----------------------------- PDF -----------------------------
def test_pdf_page_numbers(tmp_path: Path):
    sample = FIXTURES / "sample.pdf"
    sections = PdfParser().parse(sample)
    assert len(sections) == 3
    assert [s.page_number for s in sections] == [1, 2, 3]
    assert "Page One" in sections[0].content
    assert "Page Two" in sections[1].content
    assert "Page Three" in sections[2].content
    # PDF 不设置 heading / paragraph_number / line_start
    assert sections[0].heading is None
    assert sections[0].paragraph_number is None
    assert sections[0].line_start is None


def test_pdf_scanned_raises(tmp_path: Path):
    scanned = FIXTURES / "scanned.pdf"
    with pytest.raises(UnsupportedScannedPDFError):
        PdfParser().parse(scanned)


# ----------------------------- 统一入口 -----------------------------
def test_parse_document_routes_by_extension(tmp_path: Path):
    """parse_document 按扩展名分发。"""
    txt = tmp_path / "x.txt"
    txt.write_text("测试内容\n", encoding="utf-8")
    sections = parse_document(txt)
    assert isinstance(sections[0], RawDocumentSection)
    assert sections[0].source_name == "x.txt"


def test_parse_document_unsupported_extension(tmp_path: Path):
    f = tmp_path / "x.xyz"
    f.write_text("nope", encoding="utf-8")
    with pytest.raises(UnsupportedFileTypeError):
        parse_document(f)


def test_parse_document_not_found(tmp_path: Path):
    with pytest.raises(FileNotFoundError):
        parse_document(tmp_path / "no_such.pdf")


def test_supported_extensions_includes_required():
    exts = supported_extensions()
    for e in (".pdf", ".docx", ".txt", ".md"):
        assert e in exts


# ===================== 阶段 2.1：DOCX 标题启发式识别 =====================
def test_docx_standard_heading_style_recognized(tmp_path: Path):
    """第一优先级：Word Heading 样式被识别为标题。"""
    from docx import Document

    doc = Document()
    doc.add_heading("第一章 总则", level=1)
    doc.add_paragraph("这是第一章的正文内容。")
    doc.add_heading("1.1 子标题", level=2)
    doc.add_paragraph("这是子章节的正文内容。")
    out = tmp_path / "heading.docx"
    doc.save(str(out))

    sections = DocxParser().parse(out)
    # 第一段正文继承 "第一章 总则"
    assert sections[0].heading == "第一章 总则"
    assert sections[1].heading == "1.1 子标题"


def test_docx_heuristic_heading_recognized_via_pattern(tmp_path: Path):
    """第二优先级：匹配常见标题模式（如"一、"）且满足多信号 → 启发式识别。"""
    from docx import Document

    doc = Document()
    doc.add_paragraph("本目录说明这是一段较长的开头正文，用于建立基线字号。")
    doc.add_paragraph("一、概述")  # 短 + pattern → strong_signals = 2
    doc.add_paragraph("概述段内容。")
    out = tmp_path / "heur.docx"
    doc.save(str(out))

    sections = DocxParser().parse(out)
    # 第一段继承 None（开头无标题）；第二段继承 "一、概述"
    assert sections[0].heading is None
    assert sections[1].heading == "一、概述"


def test_docx_short_drug_name_not_misclassified_as_heading(tmp_path: Path):
    """普通短药品名称（无 pattern、无加粗、无大字号）不会被误判为标题。"""
    from docx import Document

    doc = Document()
    doc.add_paragraph("本节描述化学药品的分类与目录。")
    # 这些短段落只触发 "short" 一个 strong signal → 不算标题
    doc.add_paragraph("青霉素")
    doc.add_paragraph("头孢菌素")
    doc.add_paragraph("阿莫西林")
    out = tmp_path / "drugs.docx"
    doc.save(str(out))

    sections = DocxParser().parse(out)
    # 4 个段落都保留（heading 为 None）
    assert len(sections) == 4
    for s in sections:
        assert s.heading is None, f"误判为标题: {s.content!r} (heading={s.heading!r})"


def test_docx_long_paragraph_not_misclassified_as_heading(tmp_path: Path):
    """过长段落即使字号大也不会被识别为标题。"""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    # 长段落（大字号但超过标题长度阈值）
    p = doc.add_paragraph("这是一段非常长的内容，" * 10)
    for run in p.runs:
        run.font.size = Pt(20)
    out = tmp_path / "long.docx"
    doc.save(str(out))

    sections = DocxParser().parse(out)
    # 长段落被识别为正文，不是标题
    assert len(sections) == 1
    assert sections[0].heading is None


def test_docx_normal_drug_name_still_sections(tmp_path: Path):
    """含中文目录编号 + 长正文 + 短药品名时，正文继承最近标题。"""
    from docx import Document

    doc = Document()
    doc.add_paragraph("这是一段用于建立基线的较长开头正文。")
    doc.add_paragraph("一、化学药品")  # 启发式识别为标题
    doc.add_paragraph("青霉素")
    doc.add_paragraph("头孢菌素")
    out = tmp_path / "drugs2.docx"
    doc.save(str(out))

    sections = DocxParser().parse(out)
    # 第一段继承 None；第二、三段继承 "一、化学药品"
    assert sections[0].heading is None
    assert sections[1].heading == "一、化学药品"
    assert sections[2].heading == "一、化学药品"


def test_docx_does_not_split_long_pattern_only():
    """纯 pattern 信号 + 长文本不被识别（需要至少 2 个 strong signal）。"""
    from docx import Document
    from rag.parsers import DocxParser

    doc = Document()
    # "1. " 开头但很长（> 30 字符）→ 只触发 pattern → strong_signals < 2
    doc.add_paragraph("1. " + "本段落是长正文，" * 10)
    out = __import__("pathlib").Path(__file__).parent / "_tmp_long_pattern.docx"
    doc.save(str(out))
    try:
        sections = DocxParser().parse(out)
        assert len(sections) == 1
        assert sections[0].heading is None
    finally:
        out.unlink(missing_ok=True)