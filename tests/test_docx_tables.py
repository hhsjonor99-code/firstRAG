"""DOCX 表格解析测试（阶段 2.2）。

覆盖：
- 表格基本提取（表头识别 + 结构化行）
- 段落与表格交错顺序保持
- 表格继承最近标题
- 不同表格之间不合并
- 合并单元格去重
- 表格行元数据 (table_index / row_start / row_end / column_names)
- 短药品名不被误判为标题
- chunk_index 连续、无空 chunk
- 元数据在分块后保留
"""

from __future__ import annotations

import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from rag.models import DocumentChunk  # noqa: E402
from rag.parsers import DocxParser, RawDocumentSection, parse_document  # noqa: E402
from rag.splitter import SplitOptions, split_sections  # noqa: E402


def _make_docx_with_table(tmp_path: Path, name: str = "test.docx"):
    """构造一个含标题 + 表格 + 段落的 DOCX。

    表格 1：3 行（1 表头 + 2 数据） × 3 列
    表格 2：3 行（1 表头 + 2 数据） × 3 列
    """
    from docx import Document

    doc = Document()
    doc.add_heading("化学药品部分", level=1)
    doc.add_paragraph("本目录收录化学药品与生物制品。")
    # 表格 1：含表头 + 2 行数据
    t1 = doc.add_table(rows=3, cols=3)
    t1.style = "Table Grid"
    t1.cell(0, 0).text = "序号"
    t1.cell(0, 1).text = "药品名称"
    t1.cell(0, 2).text = "剂型"
    t1.cell(1, 0).text = "1"
    t1.cell(1, 1).text = "青霉素"
    t1.cell(1, 2).text = "注射剂"
    t1.cell(2, 0).text = "2"
    t1.cell(2, 1).text = "阿莫西林"
    t1.cell(2, 2).text = "片剂"
    doc.add_paragraph("中成药部分。")
    # 表格 2：含表头 + 2 行数据
    t2 = doc.add_table(rows=3, cols=3)
    t2.style = "Table Grid"
    t2.cell(0, 0).text = "序号"
    t2.cell(0, 1).text = "名称"
    t2.cell(0, 2).text = "规格"
    t2.cell(1, 0).text = "3"
    t2.cell(1, 1).text = "板蓝根颗粒"
    t2.cell(1, 2).text = "10g"
    t2.cell(2, 0).text = "4"
    t2.cell(2, 1).text = "感冒清热颗粒"
    t2.cell(2, 2).text = "12g"
    out = tmp_path / name
    doc.save(str(out))
    return out


def _make_docx_with_table_no_header(tmp_path: Path):
    """构造一个无表头、仅数据的表格。

    使用长文本首行避免被识别为表头。
    """
    from docx import Document

    doc = Document()
    doc.add_paragraph("数据表如下：")
    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    # 首行使用长文本（>30 字符）→ 不被识别为表头
    t.cell(0, 0).text = "这是较长的药品名称说明列，用以防止识别为表头。"
    t.cell(0, 1).text = "规格说明列。"
    t.cell(1, 0).text = "青霉素"
    t.cell(1, 1).text = "0.5g"
    out = tmp_path / "noheader.docx"
    doc.save(str(out))
    return out


# ----------------------------------------------------------------------
# 1. 表格提取基本
# ----------------------------------------------------------------------
def test_docx_extracts_table_rows(tmp_path: Path):
    """DOCX 表格行被提取为独立 section（每行一个 section）。"""
    docx = _make_docx_with_table(tmp_path)
    sections = DocxParser().parse(docx)
    # 表格 1: 3 行 (1 表头 + 2 数据) → 2 个 section
    # 表格 2: 3 行 (1 表头 + 2 数据) → 2 个 section
    # 期望 4 个表格行 section
    table_sections = [s for s in sections if s.block_type == "table"]
    assert len(table_sections) == 4, f"期望 4 个表格行，实际 {len(table_sections)}"
    # 检查结构化文本
    assert "药品名称: 青霉素" in table_sections[0].content
    assert "剂型: 注射剂" in table_sections[0].content


def test_docx_table_uses_column_names_from_header(tmp_path: Path):
    """有表头时使用 ``列名: 值`` 格式。"""
    docx = _make_docx_with_table(tmp_path)
    sections = DocxParser().parse(docx)
    table_sections = [s for s in sections if s.block_type == "table"]
    # 第 1 行: 序号: 1; 药品名称: 青霉素; 剂型: 注射剂
    first_row = table_sections[0].content
    assert "序号: 1" in first_row
    assert "药品名称: 青霉素" in first_row
    assert "剂型: 注射剂" in first_row


def test_docx_table_uses_index_when_no_header(tmp_path: Path):
    """无表头时使用 ``第N列: 值`` 格式。

    测试设计：首行较长（>30 字符）→ 不被识别为表头；或首行含终止标点。
    """
    from docx import Document

    doc = Document()
    doc.add_paragraph("数据表如下：")
    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    # 首行使用长文本 + 终止标点 → 不会被识别为表头
    t.cell(0, 0).text = "这是较长的药品名称列，含终止标点。"
    t.cell(0, 1).text = "这是规格列。"
    t.cell(1, 0).text = "青霉素"
    t.cell(1, 1).text = "0.5g"
    out = tmp_path / "noheader2.docx"
    doc.save(str(out))

    sections = DocxParser().parse(out)
    table_sections = [s for s in sections if s.block_type == "table"]
    assert len(table_sections) == 2
    # 使用 第N列 格式
    assert "第1列: 这是较长的药品名称列" in table_sections[0].content
    assert "第1列: 青霉素" in table_sections[1].content


# ----------------------------------------------------------------------
# 2. 顺序与继承
# ----------------------------------------------------------------------
def test_docx_preserves_block_order(tmp_path: Path):
    """段落 / 表格在文档中按原始顺序输出（block_index 单调递增）。"""
    docx = _make_docx_with_table(tmp_path)
    sections = DocxParser().parse(docx)
    indices = [s.block_index for s in sections]
    assert indices == sorted(indices), f"block_index 未递增: {indices}"
    assert indices == list(range(len(sections))), "block_index 应连续"


def test_docx_table_inherits_nearest_heading(tmp_path: Path):
    """表格继承之前最近识别出的标题。"""
    docx = _make_docx_with_table(tmp_path)
    sections = DocxParser().parse(docx)
    table_sections = [s for s in sections if s.block_type == "table"]
    # 表格 1 之前的标题是 "化学药品部分"
    for ts in table_sections[:3]:
        assert ts.heading == "化学药品部分"
    # 表格 2 之前的标题是 "化学药品部分"（继承最近）
    for ts in table_sections[3:]:
        assert ts.heading == "化学药品部分"


def test_docx_does_not_merge_different_tables_in_splitter(tmp_path: Path):
    """不同表格之间不直接合并（chunk 的 table_index 列表只含单表）。"""
    docx = _make_docx_with_table(tmp_path)
    sections = DocxParser().parse(docx)
    chunks = split_sections(
        sections,
        document_id="tbl-merge",
        options=SplitOptions(chunk_size=100, chunk_overlap=0),
    )
    # 每个 chunk 不应同时包含 table_index=0 和 table_index=1
    for c in chunks:
        if c.block_type == "table":
            assert c.table_index is not None
            assert len(c.table_indices) <= 1, (
                f"发现跨表合并 chunk: table_indices={c.table_indices}"
            )


# ----------------------------------------------------------------------
# 3. 元数据
# ----------------------------------------------------------------------
def test_docx_table_section_metadata(tmp_path: Path):
    """表格 section 的 metadata / 字段正确。"""
    docx = _make_docx_with_table(tmp_path)
    sections = DocxParser().parse(docx)
    table_sections = [s for s in sections if s.block_type == "table"]
    # 表格 1 (table_index=0): 2 行
    assert table_sections[0].table_index == 0
    assert table_sections[0].row_start == 1
    assert table_sections[0].row_end == 1
    assert table_sections[0].column_names == ["序号", "药品名称", "剂型"]
    # 表格 2 (table_index=1): 2 行
    assert table_sections[2].table_index == 1
    assert table_sections[2].row_start == 1
    assert table_sections[2].row_end == 1
    assert table_sections[2].column_names == ["序号", "名称", "规格"]
    assert table_sections[3].row_start == 2


def test_docx_table_chunk_preserves_metadata(tmp_path: Path):
    """表格 chunk 保留 table_index / row_start / row_end / column_names。"""
    docx = _make_docx_with_table(tmp_path)
    sections = DocxParser().parse(docx)
    chunks = split_sections(
        sections,
        document_id="tbl-meta",
        options=SplitOptions(chunk_size=200, chunk_overlap=0),
    )
    table_chunks = [c for c in chunks if c.block_type == "table"]
    assert len(table_chunks) >= 1
    c = table_chunks[0]
    assert c.table_index is not None
    assert c.row_start is not None
    assert c.row_end is not None
    assert c.column_names is not None
    # 至少一个 chunk 含 row_start=1
    assert any(c.row_start == 1 for c in table_chunks)


def test_docx_paragraph_chunk_keeps_paragraph_range(tmp_path: Path):
    """段落 chunk 保留 paragraph_start / paragraph_end。"""
    docx = _make_docx_with_table(tmp_path)
    sections = DocxParser().parse(docx)
    # 使用较小的 chunk_size 强制切分段落
    chunks = split_sections(
        sections, document_id="para-meta", options=SplitOptions(chunk_size=20, chunk_overlap=0)
    )
    para_chunks = [c for c in chunks if c.block_type == "paragraph"]
    assert len(para_chunks) >= 1
    for c in para_chunks:
        assert c.paragraph_start is not None
        assert c.paragraph_end is not None


# ----------------------------------------------------------------------
# 4. 解析不变量
# ----------------------------------------------------------------------
def test_docx_no_empty_chunks(tmp_path: Path):
    """分块后无空 chunk。"""
    docx = _make_docx_with_table(tmp_path)
    sections = DocxParser().parse(docx)
    chunks = split_sections(sections, document_id="no-empty", options=SplitOptions(chunk_size=400))
    assert all(c.content.strip() for c in chunks)


def test_docx_chunk_index_continuous_after_table_parse(tmp_path: Path):
    """含表格的 chunk_index 仍连续递增。"""
    docx = _make_docx_with_table(tmp_path)
    sections = DocxParser().parse(docx)
    chunks = split_sections(sections, document_id="idx-tbl", options=SplitOptions(chunk_size=300))
    indices = [c.chunk_index for c in chunks]
    assert indices == list(range(len(chunks)))


def test_docx_short_drug_name_not_filtered(tmp_path: Path):
    """短药品名称不被删除（在表格中）。"""
    docx = _make_docx_with_table(tmp_path)
    sections = DocxParser().parse(docx)
    # 表格行 section 应包含 "青霉素" / "阿莫西林" / "板蓝根颗粒"
    contents = " ".join(s.content for s in sections)
    assert "青霉素" in contents
    assert "阿莫西林" in contents
    assert "板蓝根颗粒" in contents


def test_docx_filter_stats_recorded(tmp_path: Path):
    """解析完成后 ``last_filter_stats`` 包含各规则统计。"""
    docx = _make_docx_with_table(tmp_path)
    parser = DocxParser()
    parser.parse(docx)
    stats = parser.last_filter_stats
    assert "blank_dropped" in stats
    assert "table_empty_row_dropped" in stats
    assert "textbox_empty_dropped" in stats


def test_parse_document_routes_to_docx_parser(tmp_path: Path):
    """``parse_document`` 入口正确分发到 DocxParser。"""
    docx = _make_docx_with_table(tmp_path)
    sections = parse_document(docx)
    assert any(s.block_type == "table" for s in sections)


# ----------------------------------------------------------------------
# 5. 合并单元格
# ----------------------------------------------------------------------
def _make_docx_with_vmerge(tmp_path: Path) -> Path:
    """构造一个含 vMerge 的表格（手工注入 OOXML）。"""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document()
    doc.add_paragraph("含 vMerge 表格")
    t = doc.add_table(rows=3, cols=2)
    t.style = "Table Grid"
    # 表头
    t.cell(0, 0).text = "类别"
    t.cell(0, 1).text = "数量"
    # 第 1 行：A类 / 10
    t.cell(1, 0).text = "A类"
    t.cell(1, 1).text = "10"
    # 第 2 行：(1,0) 续行 → vMerge continue，文本留空
    t.cell(2, 0).text = ""
    t.cell(2, 1).text = "20"
    # 在 (2,0) 的 tc 上插入 vMerge=continue
    tc_20 = t.cell(2, 0)._tc
    vMerge_continue = tc_20.makeelement(qn("w:vMerge"), {})
    tc_20.insert(0, vMerge_continue)

    out = tmp_path / "vmerge.docx"
    doc.save(str(out))
    return out


def test_docx_vmerge_dedup(tmp_path: Path):
    """vMerge=continue 的单元格文本取自上方，不重复输出。"""
    from docx import Document

    docx = _make_docx_with_vmerge(tmp_path)
    sections = DocxParser().parse(docx)
    table_sections = [s for s in sections if s.block_type == "table"]
    # 表格 1: 3 行 (1 表头 + 2 数据) → 2 个 section
    assert len(table_sections) == 2
    # 第 1 行：A类 / 10
    assert "A类" in table_sections[0].content
    # 第 2 行：vMerge continue → "A类"（不重复输出空字符串）应被填回
    # 验证：第 2 行的"类别"列应为 "A类"，不是空
    assert "A类" in table_sections[1].content


# ----------------------------------------------------------------------
# 6. 真实文件存在性（如果存在）
# ----------------------------------------------------------------------
REAL_DOCX = ROOT / "国家基本药物目录（2026年版）(OCR).docx"


@pytest.mark.skipif(not REAL_DOCX.exists(), reason="真实示例 DOCX 不存在")
def test_real_docx_coverage_above_threshold():
    """真实示例 DOCX 覆盖率 ≥ 90%（表格内容提取完整）。"""
    sections = DocxParser().parse(REAL_DOCX)
    total_chars = sum(len(s.content) for s in sections)

    # 简易覆盖率：与 python-docx 顶层段落 + 表格字符数对比
    from docx import Document

    doc = Document(str(REAL_DOCX))
    para_chars = sum(len(p.text.strip()) for p in doc.paragraphs if p.text and p.text.strip())
    table_chars = sum(
        len(c.text.strip())
        for tbl in doc.tables
        for row in tbl.rows
        for c in row.cells
        if c.text and c.text.strip()
    )
    expected = para_chars + table_chars
    coverage = total_chars / expected * 100 if expected else 0
    assert coverage >= 90, f"覆盖率过低: {coverage:.2f}% (extracted {total_chars} / {expected})"


@pytest.mark.skipif(not REAL_DOCX.exists(), reason="真实示例 DOCX 不存在")
def test_real_docx_table_count_nonzero():
    """真实示例 DOCX 解析后至少包含一个表格块。"""
    sections = DocxParser().parse(REAL_DOCX)
    table_sections = [s for s in sections if s.block_type == "table"]
    assert len(table_sections) >= 10, f"表格块过少: {len(table_sections)}"
