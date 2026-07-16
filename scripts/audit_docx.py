"""DOCX 结构审计脚本（阶段 2.2）。

只读审计真实 DOCX 文件底层的 OOXML 结构，对比：
1. python-docx 暴露的段落 / 表格 / 文本框 / 绘图对象 / 页眉页脚中的可见文本
2. 当前 DocxParser 实际提取到的 section 字符数

输出统计信息；不输出大段原文。

运行：
    python scripts/audit_docx.py [path-to-docx]

默认使用项目根目录下的 "国家基本药物目录（2026年版）(OCR).docx"。
"""

from __future__ import annotations

import os
import sys
import zipfile
from collections import Counter
from pathlib import Path
from xml.etree import ElementTree as ET

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

DEFAULT_DOCX = ROOT / "国家基本药物目录（2026年版）(OCR).docx"

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _el_text(elem) -> str:
    """抽取一个 OOXML 元素下所有 w:t 的文本。"""
    return "".join(t.text or "" for t in elem.iter(f"{W_NS}t"))


def _para_text(p) -> str:
    """抽取一个 w:p 下所有 w:t 的文本。"""
    return "".join(t.text or "" for t in p.iter(f"{W_NS}t"))


def _is_meaningful(text: str) -> bool:
    """判断文本是否非空（去掉空白）。"""
    return bool(text and text.strip())


def audit_document_xml(docx_path: Path) -> dict:
    """统计 document.xml 中的可见结构与文本。

    重要：区分"独立段落字符"与"表格内段落字符"。
    直接子级 body 的 w:p 是独立段落；嵌在 w:tc 里的 w:p 属于表格。
    """
    with zipfile.ZipFile(docx_path, "r") as zf:
        # 1. document.xml 主体
        with zf.open("word/document.xml") as f:
            doc_root = ET.parse(f).getroot()

        body = doc_root.find(f"{W_NS}body")
        p_count = 0
        tbl_count = 0
        drawing_count = 0
        pict_count = 0
        txbx_count = 0
        instr_count = 0

        # 独立段落（body 的直接子级 w:p）
        non_empty_para = 0
        para_total_chars = 0
        para_first_text = ""
        para_max_sample = ""

        # 表格统计
        total_rows = 0
        total_cells = 0
        non_empty_cells = 0
        cell_total_chars = 0
        first_cell_sample = ""

        # 文本框 / 绘图
        txbx_chars = 0
        txbx_first = ""
        drawing_text_chars = 0
        pict_text_chars = 0

        # 区分独立段落与表格内段落
        for child in list(body):
            tag = child.tag
            if tag == f"{W_NS}p":
                p_count += 1
                txt = _para_text(child)
                if _is_meaningful(txt):
                    non_empty_para += 1
                    para_total_chars += len(txt.strip())
                    if not para_first_text:
                        para_first_text = txt.strip()[:60]
                    if len(txt.strip()) > len(para_max_sample):
                        para_max_sample = txt.strip()[:60]
            elif tag == f"{W_NS}tbl":
                tbl_count += 1
                for tr in child.iter(f"{W_NS}tr"):
                    total_rows += 1
                    for tc in tr.iter(f"{W_NS}tc"):
                        total_cells += 1
                        ctxt = _para_text(tc)
                        if _is_meaningful(ctxt):
                            non_empty_cells += 1
                            cell_total_chars += len(ctxt.strip())
                            if not first_cell_sample:
                                first_cell_sample = ctxt.strip()[:60]
            elif tag == f"{W_NS}txbxContent":
                txbx_count += 1
                t_text = _para_text(child)
                if _is_meaningful(t_text):
                    txbx_chars += len(t_text.strip())
                    if not txbx_first:
                        txbx_first = t_text.strip()[:60]

        # drawing/pict/instrText 仍按全文档统计（它们可能嵌在段落或表格）
        for d in doc_root.iter(f"{W_NS}drawing"):
            drawing_count += 1
            for t in d.iter(f"{W_NS}t"):
                if t.text:
                    drawing_text_chars += len(t.text)
        for pic in doc_root.iter(f"{W_NS}pict"):
            pict_count += 1
            for t in pic.iter(f"{W_NS}t"):
                if t.text:
                    pict_text_chars += len(t.text)
        for it in doc_root.iter(f"{W_NS}instrText"):
            instr_count += 1

        # 2. 页眉 / 页脚
        header_text_chars = 0
        header_count = 0
        footer_text_chars = 0
        footer_count = 0
        for name in zf.namelist():
            base = name.split("/")[-1].lower()
            if base.startswith("header") and name.endswith(".xml"):
                with zf.open(name) as f:
                    hroot = ET.parse(f).getroot()
                for t in hroot.iter(f"{W_NS}t"):
                    if t.text and t.text.strip():
                        header_text_chars += len(t.text.strip())
                        header_count += 1
            elif base.startswith("footer") and name.endswith(".xml"):
                with zf.open(name) as f:
                    froot = ET.parse(f).getroot()
                for t in froot.iter(f"{W_NS}t"):
                    if t.text and t.text.strip():
                        footer_text_chars += len(t.text.strip())
                        footer_count += 1

        # 3. document.xml 中唯一的可见总字符（去重）
        all_text_visible = sum(
            len(t.text.strip())
            for t in doc_root.iter(f"{W_NS}t")
            if t.text and t.text.strip()
        )

        return {
            "file_size_bytes": docx_path.stat().st_size,
            "document_xml_size_bytes": zf.getinfo("word/document.xml").file_size,
            "top_level_w_p": p_count,
            "w_tbl_total": tbl_count,
            "w_tr_total": total_rows,
            "w_tc_total": total_cells,
            "w_tc_non_empty": non_empty_cells,
            "w_tc_total_chars": cell_total_chars,
            "first_cell_sample": first_cell_sample,
            "w_drawing_total": drawing_count,
            "w_pict_total": pict_count,
            "w_txbxContent_total": txbx_count,
            "w_instrText_total": instr_count,
            "para_non_empty": non_empty_para,
            "para_total_chars": para_total_chars,
            "first_para_sample": para_first_text,
            "longest_para_sample": para_max_sample,
            "txbx_total_chars": txbx_chars,
            "txbx_first_sample": txbx_first,
            "drawing_text_chars": drawing_text_chars,
            "pict_text_chars": pict_text_chars,
            "header_text_chars": header_text_chars,
            "header_text_segments": header_count,
            "footer_text_chars": footer_text_chars,
            "footer_text_segments": footer_count,
            "all_text_visible_chars": all_text_visible,
        }


def audit_python_docx_view(docx_path: Path) -> dict:
    """通过 python-docx 暴露的接口统计段落 / 表格 / section 等。"""
    from docx import Document

    doc = Document(str(docx_path))
    para_total = len(doc.paragraphs)
    para_non_empty = 0
    para_chars = 0
    for p in doc.paragraphs:
        t = p.text or ""
        if t.strip():
            para_non_empty += 1
            para_chars += len(t.strip())

    table_total = len(doc.tables)
    tbl_total_rows = 0
    tbl_total_cells = 0
    tbl_non_empty_cells = 0
    tbl_total_chars = 0
    for tbl in doc.tables:
        for row in tbl.rows:
            tbl_total_rows += 1
            for cell in row.cells:
                tbl_total_cells += 1
                ctxt = cell.text or ""
                if ctxt.strip():
                    tbl_non_empty_cells += 1
                    tbl_total_chars += len(ctxt.strip())

    inline_shapes = len(doc.inline_shapes)
    sections_count = len(doc.sections)
    return {
        "doc_paragraphs": para_total,
        "doc_paragraphs_non_empty": para_non_empty,
        "doc_paragraph_chars": para_chars,
        "doc_tables": table_total,
        "doc_table_rows": tbl_total_rows,
        "doc_table_cells": tbl_total_cells,
        "doc_table_non_empty_cells": tbl_non_empty_cells,
        "doc_table_cell_chars": tbl_total_chars,
        "doc_inline_shapes": inline_shapes,
        "doc_sections": sections_count,
    }


def audit_current_parser(docx_path: Path) -> dict:
    """运行当前 DocxParser，记录 section 统计。"""
    from rag.parsers import DocxParser

    sections = DocxParser().parse(docx_path)
    total_chars = sum(len(s.content) for s in sections)
    headings = [s.content for s in sections if s.heading is not None]
    # paragraph_number 列表
    para_nums = sorted({s.paragraph_number for s in sections if s.paragraph_number is not None})
    return {
        "parser_section_count": len(sections),
        "parser_total_chars": total_chars,
        "parser_section_with_heading": len(headings),
        "parser_para_num_min": para_nums[0] if para_nums else None,
        "parser_para_num_max": para_nums[-1] if para_nums else None,
        "parser_unique_headings": sorted(set(headings))[:10],
        "parser_section_count_per_heading": Counter(headings),
    }


def fmt(n: int | float | None) -> str:
    if n is None:
        return "-"
    if isinstance(n, int):
        return f"{n:,}"
    return f"{n:,.2f}"


def main() -> int:
    docx_path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_DOCX
    if not docx_path.exists():
        print(f"文件不存在: {docx_path}")
        return 2

    print("=" * 72)
    print(f"DOCX 结构审计：{docx_path.name}")
    print("=" * 72)
    print(f"绝对路径：{docx_path}")
    print(f"文件大小：{fmt(docx_path.stat().st_size)} 字节")

    oox = audit_document_xml(docx_path)
    pyd = audit_python_docx_view(docx_path)
    par = audit_current_parser(docx_path)

    print("\n[1] document.xml 元素统计")
    print(f"  顶层 w:p          = {fmt(oox['top_level_w_p'])}")
    print(f"  w:tbl             = {fmt(oox['w_tbl_total'])}")
    print(f"  w:tr (总行数)     = {fmt(oox['w_tr_total'])}")
    print(f"  w:tc (总单元格)   = {fmt(oox['w_tc_total'])}")
    print(f"  w:tc (非空)       = {fmt(oox['w_tc_non_empty'])}")
    print(f"  w:drawing         = {fmt(oox['w_drawing_total'])}")
    print(f"  w:pict            = {fmt(oox['w_pict_total'])}")
    print(f"  w:txbxContent     = {fmt(oox['w_txbxContent_total'])}")
    print(f"  w:instrText       = {fmt(oox['w_instrText_total'])}")

    print("\n[2] document.xml 可见文本统计")
    print(f"  非空段落数         = {fmt(oox['para_non_empty'])} / {fmt(oox['top_level_w_p'])}")
    print(f"  段落可见字符数     = {fmt(oox['para_total_chars'])}")
    print(f"  表格可见字符数     = {fmt(oox['w_tc_total_chars'])}")
    print(f"  文本框可见字符数   = {fmt(oox['txbx_total_chars'])}")
    print(f"  drawing 内文本     = {fmt(oox['drawing_text_chars'])}")
    print(f"  pict 内文本        = {fmt(oox['pict_text_chars'])}")
    print(f"  页眉可见字符数     = {fmt(oox['header_text_chars'])} ({fmt(oox['header_text_segments'])} 段)")
    print(f"  页脚可见字符数     = {fmt(oox['footer_text_chars'])} ({fmt(oox['footer_text_segments'])} 段)")
    print(f"  全文档可见总字符数 = {fmt(oox['all_text_visible_chars'])}")
    print(f"  首段样本           = {oox['first_para_sample']!r}")
    print(f"  最长段样本         = {oox['longest_para_sample']!r}")
    if oox["first_cell_sample"]:
        print(f"  首格样本           = {oox['first_cell_sample']!r}")
    if oox["txbx_first_sample"]:
        print(f"  首文本框样本       = {oox['txbx_first_sample']!r}")

    print("\n[3] python-docx 顶层视图")
    print(f"  document.paragraphs      = {fmt(pyd['doc_paragraphs'])} (非空 {fmt(pyd['doc_paragraphs_non_empty'])}, 字符 {fmt(pyd['doc_paragraph_chars'])})")
    print(f"  document.tables          = {fmt(pyd['doc_tables'])}")
    print(f"  table 总行 / 总格        = {fmt(pyd['doc_table_rows'])} / {fmt(pyd['doc_table_cells'])}")
    print(f"  table 非空格 / 字符      = {fmt(pyd['doc_table_non_empty_cells'])} / {fmt(pyd['doc_table_cell_chars'])}")
    print(f"  document.inline_shapes   = {fmt(pyd['doc_inline_shapes'])}")
    print(f"  document.sections        = {fmt(pyd['doc_sections'])}")

    print("\n[4] 当前 DocxParser 提取")
    print(f"  有效 section 数     = {fmt(par['parser_section_count'])}")
    print(f"  提取总字符数        = {fmt(par['parser_total_chars'])}")
    print(f"  含 heading section  = {fmt(par['parser_section_with_heading'])}")
    print(f"  段落号范围          = {par['parser_para_num_min']} ~ {par['parser_para_num_max']}")
    print(f"  heading 计数（前10） = {dict(list(par['parser_section_count_per_heading'].items())[:10])}")

    coverage = 0.0
    if oox["all_text_visible_chars"] > 0:
        coverage = par["parser_total_chars"] / oox["all_text_visible_chars"] * 100
    print("\n[5] 覆盖率对比")
    print(f"  document.xml 可见总字符 = {fmt(oox['all_text_visible_chars'])}")
    print(f"  parser 提取总字符       = {fmt(par['parser_total_chars'])}")
    print(f"  覆盖率                  = {coverage:.2f}%")
    print(f"  表格字符占比            = {oox['w_tc_total_chars'] / max(oox['all_text_visible_chars'], 1) * 100:.2f}%")
    print(f"  顶层段落字符占比        = {oox['para_total_chars'] / max(oox['all_text_visible_chars'], 1) * 100:.2f}%")

    print("\n[6] 主体内容结论（自动判断）")
    # 主体判定：哪个类型占比最大
    p_chars = oox["para_total_chars"]
    t_chars = oox["w_tc_total_chars"]
    b_chars = oox["txbx_total_chars"]
    print(f"  顶层段落可见字符 : 表格可见字符 : 文本框可见字符 = "
          f"{fmt(p_chars)} : {fmt(t_chars)} : {fmt(b_chars)}")
    if t_chars > p_chars and t_chars > b_chars and t_chars > 0:
        print("  → 主体内容大量集中在 Word 表格中")
    elif p_chars > t_chars and p_chars > b_chars:
        print("  → 主体内容主要在普通段落中")
    elif b_chars > p_chars and b_chars > t_chars:
        print("  → 主体内容主要在文本框中")
    else:
        print("  → 多种结构并存")
    return 0


if __name__ == "__main__":
    sys.exit(main())
